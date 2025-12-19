"""Core utilities for working with DOCX files."""

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional
from io import BytesIO

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class ImageInfo:
    """Information about an image extracted from a document."""
    index: int
    rel_id: str
    filename: str
    content_type: str
    size_bytes: int
    width: Optional[int] = None
    height: Optional[int] = None
    temp_path: Optional[Path] = None


@dataclass
class ParagraphInfo:
    """Information about a paragraph in a document."""
    index: int
    text: str
    style: str
    is_heading: bool
    heading_level: Optional[int] = None


@dataclass
class DocumentStructure:
    """Complete structure analysis of a document."""
    total_paragraphs: int
    paragraphs: List[ParagraphInfo]
    headings: List[ParagraphInfo]
    empty_paragraphs: List[int]


def extract_images_from_docx(doc_path: Path, output_dir: Path) -> List[ImageInfo]:
    """Extract all images from a DOCX file.

    Args:
        doc_path: Path to the source DOCX file
        output_dir: Directory to save extracted images

    Returns:
        List of ImageInfo objects containing image metadata
    """
    logger.info(f"从 {doc_path} 提取图片")
    doc = Document(doc_path)
    images = []
    image_index = 0

    # Ensure output directory exists
    output_dir.mkdir(parents=True, exist_ok=True)

    # Iterate through all relationships to find images
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_data = image_part.blob
            content_type = image_part.content_type

            # Determine file extension
            ext = content_type.split('/')[-1]
            if ext == 'jpeg':
                ext = 'jpg'

            # Save image to temp directory
            filename = f"image_{image_index}.{ext}"
            temp_path = output_dir / filename

            with open(temp_path, 'wb') as f:
                f.write(image_data)

            # Get image dimensions using PIL
            try:
                with Image.open(BytesIO(image_data)) as img:
                    width, height = img.size
            except Exception as e:
                logger.warning(f"无法获取 {filename} 的尺寸: {e}")
                width, height = None, None

            # Create ImageInfo object
            img_info = ImageInfo(
                index=image_index,
                rel_id=rel_id,
                filename=filename,
                content_type=content_type,
                size_bytes=len(image_data),
                width=width,
                height=height,
                temp_path=temp_path
            )
            images.append(img_info)
            image_index += 1

            logger.debug(f"Extracted {filename} ({len(image_data)} bytes)")

    logger.info(f"提取 {len(images)} 张图片到 {output_dir}")
    return images


def analyze_document_structure(doc_path: Path) -> DocumentStructure:
    """Analyze the structure of a DOCX document.

    Args:
        doc_path: Path to the DOCX file

    Returns:
        DocumentStructure object containing document analysis
    """
    logger.info(f"分析文档结构: {doc_path}")
    doc = Document(doc_path)
    paragraphs_info = []
    headings = []
    empty_paragraphs = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        # Determine if it's a heading
        is_heading = style.startswith('Heading')
        heading_level = None

        if is_heading:
            try:
                heading_level = int(style.split()[-1])
            except (ValueError, IndexError):
                heading_level = None

        # Track empty paragraphs
        if not text:
            empty_paragraphs.append(idx)

        para_info = ParagraphInfo(
            index=idx,
            text=text,
            style=style,
            is_heading=is_heading,
            heading_level=heading_level
        )
        paragraphs_info.append(para_info)

        if is_heading:
            headings.append(para_info)

    structure = DocumentStructure(
        total_paragraphs=len(doc.paragraphs),
        paragraphs=paragraphs_info,
        headings=headings,
        empty_paragraphs=empty_paragraphs
    )

    logger.info(f"文档包含 {structure.total_paragraphs} 个段落, "
                f"{len(headings)} 个标题, {len(empty_paragraphs)} 个空段落")

    return structure


def insert_image_at_paragraph(
    doc: Document,
    para_index: int,
    image_path: Path,
    width_inches: float = 6.0,
    alignment: str = "center"
) -> None:
    """Insert an image after a specific paragraph.

    Args:
        doc: Document object
        para_index: Index of paragraph to insert after
        image_path: Path to the image file
        width_inches: Width of the image in inches
        alignment: Alignment of the image (left, center, right)
    """
    logger.debug(f"Inserting {image_path.name} after paragraph {para_index}")

    # Get or create the paragraph to insert the image
    if para_index < len(doc.paragraphs):
        # Insert a new paragraph after the specified index
        para = doc.paragraphs[para_index]
        new_para = para.insert_paragraph_before("")
        # Move the new paragraph to after the current one
        # We'll add the image to a new paragraph instead
        para_for_image = para._element.addnext(
            doc.add_paragraph()._element
        )
        # Get the paragraph object
        for p in doc.paragraphs:
            if p._element == para_for_image:
                para = p
                break
    else:
        # Add at the end
        para = doc.add_paragraph()

    # Add the image
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    # Set alignment
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }
    para.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)

    logger.debug(f"Inserted image with width={width_inches} inches, alignment={alignment}")


def insert_images_batch(
    doc_path: Path,
    output_path: Path,
    insertion_plan: List[dict],
    width_inches: float = 6.0,
    alignment: str = "center"
) -> None:
    """Insert multiple images into a document according to a plan.

    Args:
        doc_path: Path to the target DOCX file
        output_path: Path to save the modified document
        insertion_plan: List of dicts with keys: image_path, insert_after_para
        width_inches: Width of images in inches
        alignment: Alignment of images
    """
    logger.info(f"向 {doc_path} 插入 {len(insertion_plan)} 张图片")
    doc = Document(doc_path)

    # Sort insertion plan by paragraph index in reverse order
    # This ensures that inserting images doesn't shift the indices
    sorted_plan = sorted(insertion_plan, key=lambda x: x['insert_after_para'], reverse=True)

    for plan_item in sorted_plan:
        image_path = plan_item['image_path']
        para_index = plan_item['insert_after_para']

        insert_image_at_paragraph(
            doc=doc,
            para_index=para_index,
            image_path=image_path,
            width_inches=width_inches,
            alignment=alignment
        )

    # Save the modified document
    doc.save(str(output_path))
    logger.info(f"已保存修改后的文档到 {output_path}")


def get_document_summary(structure: DocumentStructure) -> str:
    """Generate a text summary of document structure for LLM.

    Args:
        structure: DocumentStructure object

    Returns:
        Formatted string describing the document structure
    """
    lines = [f"Document Structure Summary:"]
    lines.append(f"Total Paragraphs: {structure.total_paragraphs}")
    lines.append(f"Total Headings: {len(structure.headings)}")
    lines.append(f"Empty Paragraphs: {len(structure.empty_paragraphs)}")
    lines.append("")
    lines.append("Paragraph Details:")

    for para in structure.paragraphs:
        prefix = f"Para {para.index}"
        if para.is_heading:
            prefix += f" ({para.style})"
        else:
            prefix += " (Normal)"

        text_preview = para.text[:50] if para.text else "[empty]"
        lines.append(f"  - {prefix}: {text_preview}")

    return "\n".join(lines)


if __name__ == "__main__":
    # Test the utilities
    from pathlib import Path
    import config

    print("Testing docx_utils...")

    # Test document analysis
    if config.TARGET_DOC_PATH.exists():
        structure = analyze_document_structure(config.TARGET_DOC_PATH)
        print(get_document_summary(structure))
